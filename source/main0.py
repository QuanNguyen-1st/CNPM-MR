import docx
import pandas as pd
import re
my_path = "E:/Study/CNPM/MR/Word/Tu dien Ba Na - Việt Kon Tum.docx"
file = docx.Document(my_path)
language0 = []
language1 = []
para = file.paragraphs
for i in range(96, 17126):
    try:
        index = para[i].text.find('x:')
        if index == -1:
            pass
        else:
            para[i].text = para[i].text[:index]
    except IndexError:
        pass
    try:
        x, y = file.paragraphs[i].text.split(':', 1)
        if (x == "" or y == ""):
            continue
        language0.append(x.strip())
        language1.append(y.strip())
    except:
        continue

word_type = []
for i, word in enumerate(language0):
    try:
        if word[-1] == ')':
            if word[-5] == '(':
                word_type.append(word[-4:-1])
            else:
                word_type.append(word[-3:-1])
        else:
            word_type.append("")
    except:
        word_type.append("")
    try:
        language0[i] = re.sub("[\(\[].*?[\)\]]", "", language0[i])
        language0[i] = re.sub('\^', "", language0[i])
        language0[i] = re.sub('®', "", language0[i])
        language0[i] = re.sub('\(', "", language0[i])
        language0[i] = re.sub('\[', "", language0[i])
        language0[i] = re.sub('\)', "", language0[i])
        language0[i] = re.sub('\]', "", language0[i])
        language0[i] = ''.join([x for x in language0[i] if not x.isdigit()])
        language0[i] = language0[i].strip()
    except:
        pass

path = "Excel/df4.xlsx"
print(len(language0))
print(len(language1))
print(len(word_type))
data = {'language0': language0,
        'language1': language1,
        'word_type': word_type}            
df = pd.DataFrame(data)
df.to_excel(path, index = False)